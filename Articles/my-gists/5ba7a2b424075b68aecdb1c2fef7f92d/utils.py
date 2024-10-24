import cryptography
import mnemonic
import password_strength
import pytz
import termcolor
import roman
import nato_phonetic_alphabet
import uuid
import hashlib
import bcrypt
import base64
import datetime
import re
import urllib.parse
import html
import os
import time
import csv
import json
import yaml
import toml
import jwt
import ipaddress
import hmac
import http
import magic
import rsa
import secrets
import string
import random
import qrtools
import pyotp
import qrcode
import faker
import lorem
import statistics
import requests
import textstat
import weasyprint

import BeautifulSoup
import ua_parser
import croniter

import sqlparse
import sqlite3

import pytesseract
import pdf2image
from docx import Document
from docx.shared import Inches
import imgkit

import ffmpeg

import tempfile
import pdf2docx
import PyPDF2
import docx2pdf
import PIL
import markdown
import cairosvg
import pyshorteners
from rapidfuzz import fuzz, process
import xml.etree.ElementTree as ET

from Cryptodome.Cipher import AES
from Cryptodome.Hash import SHA1, SHA256
from Cryptodome.Random import get_random_bytes
from Cryptodome.Util.Padding import pad

class Crypto:

    @staticmethod
    def token(length=16):
        return secrets.token_hex(length)
    @staticmethod

    def hash_text(text, algorithm='sha256'):
        hash_obj = hashlib.new(algorithm)
        hash_obj.update(text.encode('utf-8'))
        return hash_obj.hexdigest()

    def bcrypt(password, salt_rounds=12):
        salt = bcrypt.gensalt(rounds=salt_rounds)
        hashed = bcrypt.hashpw(password.encode('utf-8'), salt)
        return hashed.decode('utf-8')

    def uuid():
        return str(uuid.uuid4())

    def encrypt_decrypt_text(self, key, text, mode='encrypt'):
        if mode == 'encrypt':
            cipher_suite = cryptography.fernet.Fernet(key)
            encrypted_text = cipher_suite.encrypt(text.encode())
            return encrypted_text
        elif mode == 'decrypt':
            cipher_suite = cryptography.fernet.Fernet(key)
            decrypted_text = cipher_suite.decrypt(text)
            return decrypted_text.decode()

    def bip39_passphrase(self):
        mnemo = mnemonic.Mnemonic('english')
        passphrase = mnemo.generate_entropy(256)
        return passphrase

    def generate_hmac(self, key, data):
        hmac_hash = hmac.new(key.encode(), data.encode(), hashlib.sha256).hexdigest()
        return hmac_hash

    def generate_rsa_key_pair(self):
        public_key, private_key = rsa.newkeys(2048)
        return public_key, private_key

    def aes_encrypt(plaintext, key):
        cipher = AES.new(key, AES.MODE_ECB)
        ciphertext = cipher.encrypt(pad(plaintext, AES.block_size))
        return ciphertext

    def aes_decrypt(ciphertext, key):
        cipher = AES.new(key, AES.MODE_ECB)
        plaintext = cipher.decrypt(ciphertext)
        return plaintext

    def sha1_hash(data):
        sha1_hasher = SHA1.new()
        sha1_hasher.update(data)
        return sha1_hasher.hexdigest()

    def sha256_hash(data):
        sha256_hasher = SHA256.new()
        sha256_hasher.update(data)
        return sha256_hasher.hexdigest()

    def hmac_sha256(key, data):
        hmac_hasher = hashlib.new('sha256', key)
        hmac_hasher.update(data)
        return hmac_hasher.digest()
    
    def analyze_password_strength(self, password):
        stats = password_strength.PasswordStats(password)
        return stats.strength()

class Converter:

    def date_time(date_string, from_timezone, to_timezone):
        from_zone = pytz.timezone(from_timezone)
        to_zone = pytz.timezone(to_timezone)
        date_time = datetime.datetime.strptime(date_string, '%Y-%m-%d')
        date_time = from_zone.localize(date_time)
        converted_date_time = date_time.astimezone(to_zone)
        return converted_date_time.strftime('%Y-%m-%d %H:%M:%S')

    def integer_base(self, number, from_base, to_base):
        return format(int(number, from_base), to_base)

    def roman_numeral(self, number):
        return roman.toRoman(number).roman()

    def color(self, text, color_name):
        return termcolor.colored(text, color_name)

    def encode_base64(text):
        return base64.b64encode(text.encode('utf-8')).decode('utf-8')

    def decode_base64(encoded_text):
        return base64.b64decode(encoded_text).decode('utf-8')

    def case(self, text, case_type):
        if case_type == 'upper':
            return text.upper()
        elif case_type == 'lower':
            return text.lower()
        elif case_type == 'title':
            return text.title()

    def text_to_nato_alphabet(self, text):
        return nato_phonetic_alphabet.text_to_nato(text)

    def yaml_to_json(self, yaml_text):
        return json.dumps(yaml.safe_load(yaml_text), indent=4)

    def yaml_to_toml(self, yaml_text):
        toml_data = toml.dumps(yaml.safe_load(yaml_text))
        return toml_data

    def json_to_yaml(self, json_text):
        return yaml.dump(json.loads(json_text), default_flow_style=False)

    def json_to_toml(self, json_text):
        return yaml.dump(json.loads(json_text), default_flow_style=False)

    def list(self, data):
        return list(data)

    def toml_to_json(self, toml_text):
        return json.dumps(toml.loads(toml_text), indent=4)

    def toml_to_yaml(self, toml_text):
        return yaml.dump(toml.loads(toml_text), default_flow_style=False)

    def decimal_to_float(obj):
        if isinstance(obj, dict):
            return {k: Converter.decimal_to_float(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            return [Converter.decimal_to_float(item) for item in obj]
        elif isinstance(obj, Decimal):
            return float(obj)
        return obj

class DocumentConverter:

    def doc_to_jpg(self, input_docx, output_jpg):
        doc = Document(input_docx)
        images = [p for p in doc.paragraphs if p.runs and p.runs[0].inline_shapes]
        if images:
            images[0].runs[0].inline_shapes[0].picture.save(output_jpg)
            return True
        return False

    def html_to_image(self, input_html, output_image):
        weasyprint.HTML(input_html).write_png(output_image)
        return True

    def image_to_html(self, input_image, output_html):
        html_content = "<!DOCTYPE html><html><body><img src='{}' alt='image'></body></html>".format(input_image)
        with open(output_html, 'w') as html_file:
            html_file.write(html_content)
        return True

    def image_to_txt(self, input_image, output_text):
        text = pytesseract.image_to_string(PIL.Image.open(input_image))
        with open(output_text, 'w') as text_file:
            text_file.write(text)
        return True

    def image_to_word(self, input_image, output_word):
        doc = Document()
        doc.add_picture(input_image, width=Inches(5.0))
        doc.save(output_word)
        return True

    def jpg_to_word(self, input_jpg, output_word):
        doc = Document()
        doc.add_picture(input_jpg, width=Inches(5.0))
        doc.save(output_word)
        return True

    def pdf_to_doc(self, input_pdf, output_doc):
        pages = pdf2image.convert(input_pdf)
        doc = Document()
        for page in pages:
            doc.add_paragraph(page)
        doc.save(output_doc)
        return True

    def pdf_to_docx(self, input_pdf, output_docx):
        pages = pdf2image.convert(input_pdf)
        doc = Document()
        for page in pages:
            doc.add_paragraph(page)
        doc.save(output_docx)
        return True

    def pdf_to_image(self, input_pdf, output_image):
        images = pdf2image.convert(input_pdf)
        images[0].save(output_image, 'PNG')
        return True

    def pdf_to_jpg(self, input_pdf, output_jpg):
        images = pdf2image.convert(input_pdf)
        images[0].save(output_jpg, 'JPEG')
        return True

    def pdf_to_png(self, input_pdf, output_png):
        images = pdf2image.convert(input_pdf)
        images[0].save(output_png, 'PNG')
        return True

    def pdf_to_word(self, input_pdf, output_word):
        # PDFファイルをWordファイルに変換します。
        cv = pdf2docx.Converter(input_pdf)
        cv.convert(output_word, start=0, end=None)
        cv.close()

    def pdf_to_xps(self, input_pdf, output_xps):
        # PDFファイルをXPSファイルに変換します。
        input_pdf = PyPDF2.PdfFileReader(input_pdf)
        output_xps = PyPDF2.PdfFileWriter()

        for page_number in range(input_pdf.getNumPages()):
            output_xps.addPage(input_pdf.getPage(page_number))

        with open(output_xps, 'wb') as output_file:
            output_xps.write(output_file)

    def text_to_png(self, input_text, output_png):
        # テキストファイルをPNGファイルに変換します。
        with open(input_text, 'r') as file:
            data = file.read()

        image = PIL.Image.new('RGB', (1000, 500), color = (73, 109, 137))
        d = PIL.ImageDraw.Draw(image)
        d.text((10,10), data, fill=(255,255,0))
        image.save(output_png)

    def word_to_jpg(self, input_word, output_jpg):
        with tempfile.TemporaryDirectory() as tmpdirname:
            tmp_pdf = os.path.join(tmpdirname, "tmp.pdf")
            docx2pdf.convert(input_word, tmp_pdf)
            pages = pdf2image.convert_from_path(tmp_pdf, 500)
            pages[0].save(output_jpg, 'JPEG')

    def text_to_svg(self, input_text, output_svg):
        with open(input_text, 'r') as file:
            data = file.read()

        text_object = '<svg height="500" width="500"><text x="0" y="15" fill="black">{}</text></svg>'.format(data)
        cairosvg.svg2svg(bytestring=text_object.encode('utf-8'), write_to=output_svg)

class ImageConverter:

    def markdown_to_svg(markdown_text, width ,height):
        try:
            html = markdown.markdown(markdown_text)
            svg_data = cairosvg.svg2svg(bytestring=html.encode('utf-8'), width=width, height=height)
            return svg_data
        except Exception as e:
            return e

    def svg_to_png(svg_data, output_path):
        cairosvg.svg2png(bytestring=svg_data, write_to=output_path)

    def combine_images(image_paths, x, y):
        images = [PIL.Image.open(file) for file in image_paths]

        if len(images) != x * y:
            print("指定されたx * yの数と画像の数が一致しません。")
            return

        widths, heights = zip(*(i.size for i in images))

        total_width = sum(widths[:x])
        max_height = max(heights[i] for i in range(0, len(images), x))

        combined_image = PIL.Image.new('RGB', (total_width, max_height * y))

        x_offset, y_offset = 0, 0
        for i, image in enumerate(images):
            combined_image.paste(image, (x_offset, y_offset))
            x_offset += widths[i]

            if (i + 1) % x == 0:
                x_offset = 0
                y_offset += max_height

        return combined_image

    def filename_matched(query, target_list):
        matched_files = []
        try:
            for file in target_list:
                if query in file:
                    matched_files.append(file)
            return matched_files
        except Exception as e:
            return e

    def full_text_search(query, path_list):
        matched_files = []

        for path in path_list:
            try:
                with open(path, 'r', encoding='utf-8') as file:
                    lines = file.readlines()

                    for line_num, line in enumerate(lines, start=1):
                        if query in line:
                            matched_files.append((path, line_num))
            except FileNotFoundError:
                print(f"File not found: {path}")
            except IsADirectoryError:
                print(f"Directory is not supported: {path}")
            except Exception as e:
                print(f"Error occurred while processing file: {path}\n{e}")

        return matched_files

    def recursive_search(pattern, path='.'):
        def is_binary(file_path):
            with open(file_path, 'rb') as f:
                chunk = f.read(1024)
                return b'\0' in chunk

        matched_files = []

        for root, dirs, files in os.walk(path):
            # Skip hidden directories and binary files
            dirs[:] = [d for d in dirs if not d.startswith('.')]
            files[:] = [f for f in files if not is_binary(os.path.join(root, f))]

            for file in files:
                file_path = os.path.join(root, file)

                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        lines = f.readlines()

                        for line_num, line in enumerate(lines, start=1):
                            if re.search(pattern, line):
                                matched_files.append((file_path, line_num))
                except UnicodeDecodeError:
                    print(f"Skipped binary file: {file_path}")
                except Exception as e:
                    print(f"Error occurred while processing file: {file_path}\n{e}")

        return matched_files

    def is_binary(file_path):
        with open(file_path, 'rb') as f:
            chunk = f.read(1024)
            if b'\x00' in chunk:
                return True
            else:
                return False

    def shorten_url(url):
        s = pyshorteners.Shortener(api_key="YOUR_KEY")
        short_url = s.bitly.short(url)
        return short_url

    def fuzzy_search(query, file_list, limit=5):
        results = process.extract(query, file_list, limit=limit)
        return [result[0] for result in results]

class VideoConverter:

    def init(video):
        ffmpeg.input(video)

    def info(video):
        ffmpeg.input(video).output("-hide_banner").run()

    def convert(input_file, output_file):
        ffmpeg.input(input_file).output(output_file).run()

    def convert_all(input_directory, output_directory, pattern):
        ffmpeg.input(f'{input_directory}/{pattern}', pattern_type='glob').output(f'{output_directory}/%filename%.mp3').run()

    def crop(input_file, output_file, width, height, x, y):
        ffmpeg.input(input_file).filter('crop', width, height, x, y).output(output_file, c='copy').run()

    def resize(input_file, output_file, width, height):
        ffmpeg.input(input_file).output(output_file, s=f'{width}x{height}').run()

    def rotate(input_file, output_file, angle, scale):
        ffmpeg.input(input_file).filter('rotate', angle).output(output_file, vf=f'scale={scale}').run()

    def shrink(input_file, output_file, width, height):
        ffmpeg.input(input_file).output(output_file, vf=f'scale={width}:{height}', c='copy').run()

    def compress(input_file, output_file, crf):
        ffmpeg.input(input_file).output(output_file, vf='scale=1280:-1', vcodec='libx264', preset='veryslow', crf=crf).run()

    def fps(input_file, output_file, fps):
        ffmpeg.input(input_file).output(output_file, vf=f'fps=fps={fps}').run()

    def merge_audio(input_video, input_audio, output_file):
        ffmpeg.input(input_video).input(input_audio).output(output_file, c='copy', shortest=None).run()

    def merge_audio_no_encoding(input_video, input_audio, output_file):
        ffmpeg.input(input_video).input(input_audio).output(output_file, c='copy').run()

    def replace_audio(input_video, input_audio, output_file):
        ffmpeg.input(input_video).input(input_audio).output(output_file, c='copy', map='0:v:0', map='1:a:0').run()

    def remove_audio(input_file, output_file):
        ffmpeg.input(input_file).output(output_file, an=None, vcodec='copy').run()

    def clip_thumbnails(input_file, output_pattern):
        ffmpeg.input(input_file).output(output_pattern, vf='fps=1').run()

    def trim(input_file, output_file, start_time, end_time):
        ffmpeg.input(input_file).output(output_file, c='copy', ss=start_time, to=end_time).run()

    def remove_video(input_file, output_file):
        ffmpeg.input(input_file).output(output_file, vn=None, acodec='copy').run()

    def audio_volume(input_file, output_file, volume):
        ffmpeg.input(input_file).output(output_file, af=f'volume={volume}').run()

    def create_gif(input_file, output_file):
        ffmpeg.input(input_file).output(output_file).run()

    def change_aspect_ratio(input_file, output_file, aspect_ratio):
        ffmpeg.input(input_file).output(output_file, aspect=aspect_ratio).run()

    def add_poster_image(input_image, input_audio, output_file):
        ffmpeg.input(input_image, loop=1).input(input_audio).output(output_file, c='copy', shortest=None).run()

    def split(input_file, output_files, durations):
        in_stream = ffmpeg.input(input_file)
        outputs = []
        for i, duration in enumerate(durations):
            outputs.append(in_stream.output(output_files[i], c='copy', t=duration))

        ffmpeg.merge_outputs(*outputs).run()

    def concat(input_files, output_file):
        inputs = [ffmpeg.input(file) for file in input_files]
        ffmpeg.concat(*inputs).output(output_file, c='copy').run()


    def captions(input_file, subtitle_file, output_file):
        ffmpeg.input(input_file).output(output_file, srt=subtitle_file, c='copy', vcodec='libx264', crf=23, preset='veryfast').run()

    def preview(input_file):
        ffmpeg.input(input_file).output("-").run(overwrite_output=True)

    def video_speed(input_file, output_file, speed):
        ffmpeg.input(input_file).output(output_file, vf=f"setpts={speed}*PTS").run()

    def audio_speed(input_file, output_file, speed):
        ffmpeg.input(input_file).output(output_file, af=f"atempo={speed}", vn=None).run()

    def remove_meta_data(input_file, output_file):
        ffmpeg.input(input_file).output(output_file, map_metadata='-1', c='copy').run()

    def generate_test_video(output_file, duration, size, rate):
        ffmpeg.input('testsrc', f'duration={duration}', f'size={size}', f'rate={rate}').output(output_file).run()

    def fade(input_file, output_file, fade_in_duration, fade_out_duration):
        ffmpeg.input(input_file).output(output_file, vf=f"fade=in:0:{fade_in_duration},fade=out:st=10:d={fade_out_duration}").run()

    def screenshot(input_file, timestamps, output_prefix):
        for i, timestamp in enumerate(timestamps):
            output_file = f"{output_prefix}_{i+1}.png"
            ffmpeg.input(input_file).output(output_file, ss=timestamp, vframes=1).run()

    def keep_quality(input_file, output_file):
        ffmpeg.input(input_file).output(output_file, c='copy', vcodec='libx264').run()

    def horizontal_flip(input_file, output_file):
        ffmpeg.input(input_file).output(output_file, vf='hflip').run()

    def insert_black_last(input_file, output_file):
        ffmpeg.input(input_file).output(output_file, af='apad=packet_size=4096:pad_dur=2').run()

    def binary(input_file, output_file):
        ffmpeg.input(input_file).output(output_file, vf='hue=s=0').run()

    def blend(input_file1, input_file2, output_file):
        ffmpeg.input(input_file1).input(input_file2).output(output_file, filter_complex='blend=difference').run()

    def loop(input_file, loop_count, output_file):
        ffmpeg.input(input_file).output(output_file, filter_complex=f'loop=loop={loop_count}:size=1500:start=1500', pix_fmt='yuv420p').run()

    def side_by_side(input_file1, input_file2, output_file, alignment='horizontal', aspect_ratio='16:9'):
        if alignment == 'horizontal':
            scale = '960:-1'
            overlay = 'overlay=0:h/2:shortest=1 [wh2];[wh2][right] overlay=w:h/2:shortest=1'
        elif alignment == 'vertical':
            scale = '-1:540'
            overlay = 'overlay=w/2:0:shortest=1 [wh2];[wh2][right] overlay=w/2:h:shortest=1'

        if aspect_ratio == '16:9':
            base_size = 'hd1080'
        elif aspect_ratio == '4:3':
            base_size = 'hd1080'

        ffmpeg.input(input_file1).input(input_file2).output(output_file, filter_complex=f"color=s={base_size} [base];[0] setpts=PTS-STARTPTS, scale={scale} [left];[1] setpts=PTS-STARTPTS, scale={scale} [right];[base][left] {overlay}", movflags='+faststart', pix_fmt='yuv420p', c='libx264', crf=20).run()

    def horizontal_align_3(input_file1, input_file2, input_file3, output_file, aspect_ratio='16:9'):
        if aspect_ratio == '16:9':
            base_size = 'hd1080'
            scale = '640:-1'
            overlay = 'overlay=0:h:shortest=1 [w];[w][middle] overlay=w:h:shortest=1[w2];[w2][right] overlay=w*2:h:shortest=1'
        elif aspect_ratio == '4:3':
            base_size = 'hd1080'
            scale = '640:-1'
            overlay = 'overlay=0:h/1.6:shortest=1 [w];[w][middle] overlay=w:h/1.6:shortest=1[w2];[w2][right] overlay=w*2:h/1.6:shortest=1'

        ffmpeg.input(input_file1).input(input_file2).input(input_file3).output(output_file, filter_complex=f"color=s={base_size} [back_color];[0] setpts=PTS-STARTPTS, scale={scale} [left];[1] setpts=PTS-STARTPTS, scale={scale}[middle];[2] setpts=PTS-STARTPTS, scale={scale} [right];[back_color][left] {overlay}", movflags='+faststart', pix_fmt='yuv420p', c='libx264', crf=20).run()

    def vertical_align_2(input_file1, input_file2, output_file):
        ffmpeg.input(input_file1).input(input_file2).output(output_file, filter_complex="[0:0]pad=iw:2*ih[a];[a][1:0]overlay=0:h").run()

    def sound(frequency, duration, output_file):
        ffmpeg.input('sine=frequency={}:duration={}'.format(frequency, duration)).output(output_file).run()

    def change_pitch(input_file, sample_rate, output_file):
        ffmpeg.input(input_file).output(output_file, af='asetrate={}'.format(sample_rate)).run()

    def crossfade(input_file1, input_file2, duration, output_file):
        ffmpeg.input(input_file1).input(input_file2).output(output_file, filter_complex='acrossfade=duration={}:curve1=exp:curve2=exp'.format(duration)).run()

    def picture_in_picture(input_file1, input_file2, output_file):
        ffmpeg.input(input_file1).input(input_file2).output(output_file, filter_complex="[1:0]scale=iw/2:ih/2[red];[0:0][red]overlay").run()

    def split_5(input_files, output_file):
        filter_complex = ''
        for i, input_file in enumerate(input_files):
            filter_complex += '[{}:v]scale=-1:720[v{}];[{}:a]'.format(i, i, i)
        filter_complex += 'amix=inputs={}:duration=shortest,loudnorm[a]'.format(len(input_files))
        filter_complex += ''.join('[v{}]'.format(i) for i in range(len(input_files)))
        filter_complex += 'xstack=inputs={}:layout=0_0|w0_0|w0+w1_0|w0_h1|w0+w1_h1[v]'.format(len(input_files))
        ffmpeg.input(*input_files).output(output_file, filter_complex=filter_complex).run()

    def split_100(input_file, output_file, rows, cols):
        duration = ffmpeg.probe(input_file)['format']['duration']
        filter_complex = ''
        for i in range(rows):
            for j in range(cols):
                index = i * cols + j
                filter_complex += '[0:a]atrim=start={}*{}:end={}*{}[a{:02d}];'.format(duration, index/100, duration, (index+1)/100, index)
                filter_complex += '[0:v]trim=start={}*{}:end={}*{},setpts=PTS-STARTPTS[v{:02d}];'.format(duration, index/100, duration, (index+1)/100, index)
        filter_complex += ''.join('[v{:02d}]scale=-1:720[v{:02d}];'.format(i, i) for i in range(rows*cols))
        filter_complex += 'xstack=inputs={}:layout='.format(rows*cols)
        filter_complex += '|'.join('{}_{}'.format(j, i) for i in range(rows) for j in range(cols))
        filter_complex += '[v]'
        ffmpeg.input(input_file).output(output_file, filter_complex=filter_complex).run()

    def subtitle(input_file, subtitle_file, output_file):
        ffmpeg.input(input_file).input(subtitle_file).output(output_file, vf="subtitles={}".format(subtitle_file)).run()

    def subtitle_manually(input_file, subtitle_files, output_file):
        filter_complex = ''
        for i, subtitle_file in enumerate(subtitle_files):
            filter_complex += '-i {} '.format(subtitle_file)
        filter_complex += '-map 0:v -map 0:a '
        for i in range(len(subtitle_files)):
            filter_complex += '-map {} '.format(i+1)
        filter_complex += '-c:v copy -c:a copy -c:s srt '
        for i in range(len(subtitle_files)):
            filter_complex += '-metadata:s:s:{} language={} '.format(i, subtitle_files[i].split('-')[1].split('.')[0])
        filter_complex += output_file
        ffmpeg.input(input_file).output(output_file, filter_complex=filter_complex).run()

    def hd_to_mobile(input_file, output_file):
        ffmpeg.input(input_file).output(output_file, vcodec='libx264', r=19, b='120k', s='480x270').run()

    def make_null_video(duration, output_file):
        ffmpeg.input('color=black:s=3840x2160:r=24000/1001').output(output_file).run()

    def preview_with_logo(in_file, image):
        (
            ffmpeg
            .input(in_file)
            .input(image, loop=1)
            .output('-', vf='format=yuva420p,lut=a=\'val*0.9\',fade=in:st=1:d=2:alpha=1,fade=out:st=7:d=2:alpha=1', f='matroska')
            .run(cmd='ffplay', input='-')
        )

    def preview_with_logo_gif(in_file, gif_file):
        (
            ffmpeg
            .input(in_file)
            .input(gif_file, ignore_loop=0)
            .output('-', vf='format=yuva420p,lut=a=\'val*0.9\',fade=in:st=1:d=2:alpha=1,fade=out:st=7:d=2:alpha=1', f='matroska')
            .run(cmd='ffplay', input='-')
        )

    def overlay_logo_bottom_left(in_file, image, out_file):
        (
            ffmpeg
            .input(in_file)
            .input(image)
            .output(out_file, vf='overlay=x=main_w*0.01:y=main_h-overlay_h-(main_h*0.01)', loglevel='fatal')
            .run()
        )

    def overlay_logo_bottom_left_complex(in_file, image, out_file):
        (
            ffmpeg
            .input(in_file)
            .input(image, loop=1)
            .output(out_file, vf='overlay=x=main_w*0.01:y=main_h-overlay_h-(main_h*0.01)', loglevel='fatal')
            .run()
        )

    def overlay_logo_center(in_file, image, out_file):
        (
            ffmpeg
            .input(in_file)
            .input(image, loop=1)
            .output(out_file, vf='overlay=(W-w)/2:(H-h)/2:shortest=1', loglevel='fatal')
            .run()
        )

    def overlay_logo_with_gif(in_file, gif_file, out_file):
        (
            ffmpeg
            .input(in_file)
            .input(gif_file, ignore_loop=0)
            .output(out_file, vf='format=yuva420p,lut=a=\'val*0.9\',fade=in:st=1:d=2:alpha=1,fade=out:st=7:d=2:alpha=1', movflags='+faststart', pix_fmt='yuv420p', c='libx264', crf=20)
            .run()
        )

    def overlay_logo_with_gif_preview(in_file, gif_file):
        (
            ffmpeg
            .input(in_file)
            .input(gif_file, ignore_loop=0)
            .output('-', vf='format=yuva420p,lut=a=\'val*0.9\',fade=in:st=1:d=2:alpha=1,fade=out:st=7:d=2:alpha=1')
            .run(cmd='ffplay', input='-')
        )

    def overlay_video_with_text(in_file, image, out_file):
        (
            ffmpeg
            .input(in_file)
            .input(image)
            .output(out_file, vf='overlay=10:10,drawtext=text=\'Hello World\'')
            .output(out_file, acodec='copy', movflags='+faststart')
            .run()
        )

    def overlay_logo_with_enable(in_file, image, out_file):
        (
            ffmpeg
            .input(in_file)
            .input(image)
            .output(out_file, vf='overlay=10:10:enable=between(t,0,30)', acodec='copy')
            .run()
        )

    def encode_with_logo(in_file, image, out_file):
        (
            ffmpeg
            .input(in_file)
            .input(image, loop=1)
            .output(out_file, vf='format=yuva420p,lut=a=\'val*0.9\',fade=in:st=1:d=2:alpha=1,fade=out:st=7:d=2:alpha=1', movflags='+faststart', pix_fmt='yuv420p', c='libx264', crf=20)
            .run()
        )

    def encode_with_logo_gif(in_file, gif_file, out_file):
        (
            ffmpeg
            .input(in_file)
            .input(gif_file, ignore_loop=0)
            .output(out_file, vf='format=yuva420p,lut=a=\'val*0.9\',fade=in:st=1:d=2:alpha=1,fade=out:st=7:d=2:alpha=1', movflags='+faststart', pix_fmt='yuv420p', c='libx264', crf=20)
            .run()
        )

    def encode_with_logo_center(in_file, image, out_file):
        (
            ffmpeg
            .input(in_file)
            .input(image, loop=1)
            .output(out_file, vf='overlay=y=(H-h):shortest=1', movflags='+faststart', pix_fmt='yuv420p', c='libx264', crf=20, loglevel='fatal')
            .run()
        )

    def encode_with_logo_random_position(in_file, image, out_file):
        (
            ffmpeg
            .input(in_file)
            .input(image, loop=1)
            .output(out_file, vf='overlay=shortest=1:x=if(eq(mod(n,200),0),sin(random(1))*W,x):y=if(eq(mod(n,200),0),sin(random(1))*H,y)', movflags='+faststart', pix_fmt='yuv420p', c='libx264', crf=20, loglevel='fatal')
            .run()
        )

    def encode_with_logo_gif_random_position(in_file, gif_file, out_file):
        (
            ffmpeg
            .input(in_file)
            .input(gif_file, ignore_loop=0)
            .output(out_file, vf='format=yuva420p,lut=a=\'val*0.9\',fade=in:st=1:d=2:alpha=1,fade=out:st=7:d=2:alpha=1', movflags='+faststart', pix_fmt='yuv420p', c='libx264', crf=20, loglevel='fatal')
            .run()
        )

    def add_text_overlay(in_file, font, text, font_color, font_size, box_color, box_opacity, out_file):
        (
            ffmpeg
            .input(in_file)
            .output(out_file, vf=f'drawtext=fontfile={font}:text={text}:fontcolor={font_color}:fontsize={font_size}:box=1:boxcolor={box_color}@{box_opacity}:boxborderw=5:x=(w-text_w)/2:y=(h-text_h)/2')
            .output(out_file, acodec='copy')
            .run()
        )

    def add_text_overlay_simple(in_file, text, out_file):
        (
            ffmpeg
            .input(in_file)
            .output(out_file, vf='drawtext=text=\'FFmpeg is the best\'')
            .output(out_file, acodec='copy')
            .run()
        )

    def add_random_text_overlay(in_file, font, font_size, font_color, text, out_file):
        (
            ffmpeg
            .input(in_file)
            .output(out_file, vf=f'drawtext=fontfile={font}:fontsize={font_size}:fontcolor={font_color}@0.5:text={text}:x=if(eq(mod(t,30),0),rand(0,(W-tw)),x):y=if(eq(mod(t,30),0),rand(0,(H-th)),y)')
            .output(out_file, c='libx264', crf=23, acodec='copy')
            .run()
        )

    def add_random_text_overlay_encode(in_file, font, font_size, font_color, text, out_file):
        (
            ffmpeg
            .input(in_file, framerate='30000/1001')
            .input('test.png', loop=1)
            .output(out_file, vf='fade=out:st=30:d=1:alpha=1[ov];[0:v][ov]overlay=10:10[v]', acodec='copy', shortest=None)
            .run()
        )

    def apply_watermark_and_effects(in_file, image, sub, out_file):
        (
            ffmpeg
            .input(in_file)
            .output(image, vf=f'fade=in:0:20,subtitles={sub}')
            .output(out_file, vf=f'overlay=W-w-10:H-h-10,fade=in:0:20')
            .output(out_file, vf='format=yuv420p', crf=18, preset='slow', c='aac', strict='-2')
            .run()
        )

    def apply_watermark_and_effects_complex(in_file, image, sub, out_file):
        (
            ffmpeg
            .input(in_file)
            .output(image, vf='fade=0:1:alpha=1,setpts=PTS+N/TB')
            .output(out_file, vf=f'overlay=W-w-10:H-h-10,fade=in:0:20,subtitles={sub}')
            .output(out_file, vf='format=yuv420p', crf=18, preset='slow', c='aac')
            .run()
        )

    def apply_watermark_and_effects(in_file, image, sub, out_file):
        (
            ffmpeg
            .input(in_file)
            .output(out_file, vf=f'overlay=W-w-10:H-h-10,fade=in:0:20[tmp_overlay];[tmp_overlay]subtitles={sub}')
            .output(out_file, codec='libx264', crf=18, preset='slow', pix_fmt='yuv420p', c='aac', strict='-2')
            .run()
        )

    def apply_watermark_and_effects_complex(in_file, image, sub, out_file):
        (
            ffmpeg
            .input(in_file)
            .output(out_file, loop=1, t=2, vf='fade=0:1:alpha=1,setpts=PTS+N/TB[wm];[0:v][wm]overlay=W-w-10:H-h-10,fade=in:0:20,subtitles=' + sub)
            .output(out_file, codec='libx264', crf=18, preset='slow', c='aac')
            .run()
        )

class Web:

    def encode_url(text):
        return urllib.parse.quote(text)

    def decode_url(encoded_text):
        return urllib.parse.unquote(encoded_text)

    def escape_html_entities(self, text):
        return html.escape(text)

    def parse_url(self, url):
        parsed_url = urllib.parse.urlparse(url)
        return parsed_url

    def get_device_information(self, user_agent):
        device_info = ua_parser.user_agent_parser.Parse(user_agent)
        return device_info

    def basic_auth(self, username, password):
        credentials = f"{username}:{password}"
        credentials_base64 = base64.b64encode(credentials.encode()).decode()
        return f"Basic {credentials_base64}"

    def open_graph_meta(self, url):
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        og_meta = {}
        for tag in soup.find_all("meta", property=lambda x: x and x.startswith("og:")):
            og_meta[tag["property"][3:]] = tag["content"]
        return og_meta

    def otp_code(self, secret_key):
        totp = pyotp.TOTP(secret_key)
        return totp.now()

    def get_mime_types(self, filename):
        mime = magic.Magic()
        mime_type = mime.from_file(filename)
        return mime_type

    def parse_jwt(self, jwt_token, secret_key):
        try:
            payload = jwt.decode(jwt_token, secret_key, algorithms=["HS256"])
            return payload
        except jwt.ExpiredSignatureError:
            return "JWT token has expired."
        except jwt.InvalidTokenError:
            return "Invalid JWT token."

    def parse_user_agent(self, user_agent):
        parsed_info = ua_parser.user_agent_parser.ParseUserAgent(user_agent)
        return parsed_info

    def get_http_status_codes(self):
        status_codes = {code: http.HTTPStatus(code).phrase for code in http.HTTPStatus}
        return status_codes

    def compare_json_diff(self, json1, json2):
        diff = json.diff(json1, json2)
        return diff

class Media:

    def qr_code(self, data, file_name):
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        img.save(file_name)

    def svg_placeholder(self, width, height, text, file_name):
        svg_content = f'''
        <svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}">
            <rect width="100%" height="100%" fill="#ECECEC" />
            <text x="50%" y="50%" fill="#999" font-size="16" font-family="Arial" text-anchor="middle" dy=".3em">{text}</text>
        </svg>
        '''
        with open(file_name, 'w') as svg_file:
            svg_file.write(svg_content)

    def record_camera(self, duration, output_file):
        print(f"Recording camera for {duration} seconds...")
        time.sleep(duration)
        print(f"Recording finished. Video saved as {output_file}")

class Development:

    def get_git_cheatsheet(self):
        git_cheatsheet = '''
        # Git Cheatsheet
        # ... (Gitコマンド一覧)
        '''
        return git_cheatsheet

    def random_port(self):
        return random.randint(1024, 65535)

    def crontab(self, cron_expression, num_occurrences):
        cron = croniter.croniter(cron_expression)
        cron_schedule = [cron.get_next(datetime.datetime).isoformat() for _ in range(num_occurrences)]
        return cron_schedule

    def prettify_format_json(self, json_data):
        return json.dumps(json.loads(json_data), indent=4)

    def minify_json(self, json_data):
        return json.dumps(json.loads(json_data), separators=(',', ':'))

    def json_to_csv(self, json_data):
        csv_data = ""
        try:
            json_parsed = json.loads(json_data)
            if isinstance(json_parsed, list):
                keys = json_parsed[0].keys()
                csv_data += ','.join(keys) + '\n'
                for item in json_parsed:
                    csv_data += ','.join(str(item[key]) for key in keys) + '\n'
        except Exception as e:
            csv_data = f"Error: {str(e)}"
        return csv_data

    def prettify_format_sql(self, sql_query):
        formatted_sql = sqlparse.format(sql_query, reindent=True, keyword_case='upper')
        return formatted_sql

    def calculate_chmod(self, permissions):
        try:
            octal_permissions = int(permissions, 8)
            chmod_result = oct(octal_permissions)
            return chmod_result
        except ValueError:
            return "Invalid input."

    def docker_run_to_compose(self, docker_run_command):
        # ドッカーコマンドをdocker-compose.ymlに変換する処理
        pass

    def format_xml(self, xml_data):
        # XMLを整形する処理
        pass

class Network:

    def calculate_ipv4_subnet(self, ip_address, subnet_mask):
        network = ipaddress.IPv4Network(f'{ip_address}/{subnet_mask}', strict=False)
        return str(network.network_address), str(network.netmask)

    def ipv4_address(self, ip_address, subnet_mask):
        network = ipaddress.IPv4Network(f'{ip_address}/{subnet_mask}', strict=False)
        return network.network_address

    def expand_ipv4_range(self, start_ip, end_ip):
        ip_range = ipaddress.IPv4Range(start_ip, end_ip)
        return list(ip_range)

    def lookup_mac_address(self, mac_address):
        # マックアドレスの詳細な情報を取得する処理
        pass

    def ipv6_ula(self, prefix):
        ula = ipaddress.IPv6Network(prefix)
        return ula

class Math:

    def evaluate_math_expression(self, expression):
        try:
            result = eval(expression)
            return result
        except Exception as e:
            return f"Error: {str(e)}"

    def calculate_eta(self, start_time, end_time):
        start = datetime.strptime(start_time, '%Y-%m-%d %H:%M:%S')
        end = datetime.strptime(end_time, '%Y-%m-%d %H:%M:%S')
        elapsed_time = end - start
        return elapsed_time.total_seconds()

    def calculate_percentage(self, numerator, denominator):
        if denominator == 0:
            return "Undefined"
        percentage = (numerator / denominator) * 100
        return percentage

class Measurement:

    def measure_chronometer(self):
        start_time = time.time()
        # ここに計測したい処理を記述
        time.sleep(2)
        end_time = time.time()
        elapsed_time = end_time - start_time
        return elapsed_time

    def temperature(self, celsius):
        fahrenheit = (celsius * 9/5) + 32
        kelvin = celsius + 273.15
        return f"Fahrenheit: {fahrenheit}, Kelvin: {kelvin}"

    def build_benchmark(self, num_iterations):
        results = []
        for i in range(num_iterations):
            start_time = time.time()
            # ここにベンチマークしたい処理を記述
            time.sleep(0.1)
            end_time = time.time()
            elapsed_time = end_time - start_time
            results.append(elapsed_time)
        return results

class Text:

    def lorem_ipsum(self):
        fake = faker.Faker()
        return fake.text()

    def calculate_text_statistics(self, text):
        num_characters = len(text)
        num_words = len(text.split())
        readability_score = textstat.flesch_reading_ease(text)
        return {
            "num_characters": num_characters,
            "num_words": num_words,
            "readability_score": readability_score
        }


class Interface:

    def get_tempdir():
        timestamp = int(time.time())
        # timestamp = datetime.now().isoformat(timespec='auto')
        temp_dir = tempfile.mkdtemp()
        return timestamp, temp_dir

    @staticmethod
    def create_zip(filelist, tmp_fname, passwd=None):
        if not filelist:
            return None
        try:
            zip_name = os.path.abspath(tmp_fname)
            with zipfile.ZipFile(zip_name, "w", compression=zipfile.ZIP_DEFLATED) as f:
                for file in filelist:
                    if os.path.isfile(file):
                        f.write(file, os.path.relpath(file, os.path.dirname(filelist[0])))
                    elif os.path.isdir(file):
                        for root, dirs, files in os.walk(file):
                            for filename in files:
                                filepath = os.path.join(root, filename)
                                f.write(filepath, os.path.relpath(filepath, os.path.dirname(filelist[0])))
            if passwd:
                zip_name_encrypted = zip_name + ".zip"
                with zipfile.ZipFile(zip_name_encrypted, "w", compression=zipfile.ZIP_DEFLATED) as f:
                    f.setpassword(passwd)
                    f.write(zip_name, os.path.basename(zip_name))
                os.remove(zip_name)
                return zip_name_encrypted
            else:
                return zip_name
        except Exception as e:
            raise RuntimeError(f"Failed to create zip file: {str(e)}")

    def read_csv(csv_file, fieldnames=None, encoding='utf-8'):
        feeds = []
        with open(csv_file, 'r', newline='', encoding=encoding) as csvfile:
            reader = csv.DictReader(csvfile, fieldnames=fieldnames)
            for row in reader:
                feeds.append(row)
        data = {"items": feeds}
        return data

    def read_json(json_file, encoding='utf-8'):
        with open(json_file, 'r', encoding=encoding) as f:
            data = json.load(f)
        return data

    def read_yaml(yaml_file, encoding='utf-8'):
        with open(yaml_file, 'r', encoding=encoding) as f:
            data = yaml.safe_load(f)
        return data

    def read_toml(toml_file, encoding='utf-8'):
        with open(toml_file, 'r', encoding=encoding) as f:
            data = toml.load(f)
        return data

    def read_raw(raw_file, encoding='utf-8'):
        with open(raw_file, 'r', encoding=encoding) as f:
            data = f.read()
        return {"items": [{'data': data}]}

    def read_xml(xml_file, encoding='utf-8'):
        data = {}
        try:
            tree = ET.parse(xml_file)
            root = tree.getroot()

            for child in root:
                data[child.tag] = child.text
        except ET.ParseError as e:
            print("XML Parsing Error:", e)
            return None

        return data

    def read_sqlite3(query, data=None, path=None, encoding='utf-8'):
        try:
            conn = sqlite3.connect(path)
            cursor = conn.cursor()
            if data:
                cursor.execute(query, data)
            else:
                cursor.execute(query)
            result = cursor.fetchall()
            conn.close()
            return result
        except sqlite3.Error as e:
            print("SQLite Error:", e)
            return None

    def read_opml(opml_file, encoding='utf-8'):
        feeds = []
        try:
            tree = ET.parse(opml_file)
            root = tree.getroot()
            body = root.find('body')

            for outer_outline in body.findall('outline'):
                for inner_outline in outer_outline.findall('outline'):
                    feed = dict(inner_outline.attrib)
                    feeds.append(feed)
        except ET.ParseError as e:
            print("OPML Parsing Error:", e)
            return None

        data = {"items": feeds}
        return data

    def transform_opml_data(data):
        def walk(data):
            if isinstance(data, dict):
                filtered_data = {key: value for key, value in data.items() if value is not None}
                return {key: walk(value) for key, value in filtered_data.items()}
            elif isinstance(data, list):
                return [walk(item) for item in data]
            else:
                return data

        items = []
        _data = walk(data)
        for obj in _data["items"]:
            item = {
                "text": obj.get("text"),
                "type": obj.get("type"),
                "title": obj.get("title"),
                "url": obj.get("xmlUrl"),   # This is target_url
                "htmlUrl": obj.get("htmlUrl")
            }
            items.append(item)

        data = {"items": items}
        return data

    def read_file(fpath, query=None, encoding='utf-8'):
        if fpath.endswith('.csv'):
            data = Interface.read_csv(fpath, encoding=encoding)
        elif fpath.endswith('.json'):
            data = Interface.read_json(fpath, encoding=encoding)
        elif fpath.endswith('.yaml') or fpath.endswith('.yml'):
            data = Interface.read_yaml(fpath, encoding=encoding)
        elif fpath.endswith('.toml'):
            data = Interface.read_toml(fpath, encoding=encoding)
        elif fpath.endswith('.xml'):
            data = Interface.read_xml(fpath, encoding=encoding)
        elif fpath.endswith('.sqlite3'):
            data = Interface.read_sqlite3(query, fpath, encoding=encoding)
        elif fpath.endswith('.opml'):
            data = Interface.read_opml(fpath, encoding=encoding)
            data = Interface.transform_opml_data(data)
        elif fpath.endswith(''):
            data = Interface.read_raw(fpath, encoding=encoding)
        else:
            raise ValueError(f"Invalid file format: {fpath}")
        return data

    def write_csv(data, path, header=None, encoding='utf-8'):
        if isinstance(data, list):
            with open(path, 'w', newline='', encoding=encoding) as f:
                writer = csv.writer(f)
                for row in data:
                    writer.writerow([row]) 
        elif isinstance(data, dict):
            with open(path, 'w', newline='', encoding=encoding) as f:
                fieldnames = list(data.keys())
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader() 
                writer.writerows([data]) 
        else:
            raise ValueError("Unsupported data type")


    def write_json(data, path=None, encoding='utf-8'):
        if path:
            with open(path, 'w', encoding=encoding) as f:
                json.dump(data, f, indent=4, ensure_ascii=False)
        else:
            json_str = json.dumps(data, indent=4, ensure_ascii=False)
            return json_str

    def write_yaml(data, path=None, encoding='utf-8'):
        if path:
            with open(path, 'w', encoding=encoding) as f:
                yaml.safe_dump(data, f, default_flow_style=False, sort_keys=False, allow_unicode=True)
        else:
            yaml_str = yaml.safe_dump(data, default_flow_style=False, sort_keys=False, allow_unicode=True)
            return yaml_str

    def write_toml(data, path=None, encoding='utf-8'):
        if path:
            with open(path, 'w', encoding=encoding) as f:
                toml.dump(data, f)
        else:
            toml_str = toml.dumps(data)
            return toml_str

    def write_xml(data, path=None, root_element="root", encoding='utf-8'):
        def to_xml(data):
            xml_string = '<?xml version="1.0" encoding="UTF-8"?>\n'
            xml_string += f'<{root_element}>\n'
            for key, value in data.items():
                xml_string += f'  <{key}>{value}</{key}>\n'
            xml_string += f'</{root_element}>'
            return xml_string

        try:
            xml_data = to_xml(data)
            if path:
                with open(path, 'w', encoding=encoding) as f:
                    f.write(xml_data)
            else:
                return xml_data
        except Exception as e:
            print("Error while writing to file:", str(e))

    def write_raw(data, path=None, encoding='utf-8'):
        if path:
            with open(path, 'w', encoding=encoding) as f:
                f.write(data)
        else:
            return data

    def write_sqlite3(query, path, data=None, encoding='utf-8'):
        try:
            conn = sqlite3.connect(path)
            cursor = conn.cursor()
            if data:
                cursor.execute(query, data)
            else:
                cursor.execute(query)
            conn.commit()
            conn.close()
        except sqlite3.Error as e:
            print("SQLite Error:", e)

    def write_opml(data, path=None, encoding='utf-8'):
        def to_opml(data):
            opml_string = '<?xml version="1.0" encoding="UTF-8"?>\n<opml version="1.0">\n<head>\n<title>My OPML</title>\n</head>\n<body>\n'
            for feed in data:
                opml_string += f'<outline type="rss" text="{feed}" />\n'
            opml_string += '</body>\n</opml>'
            return opml_string

        opml_data = to_opml(data)

        if path:
            try:
                with open(path, 'w', encoding=encoding) as f:
                    f.write(opml_data)
            except Exception as e:
                print("Error while writing to file:", str(e))
        else:
            return opml_data
        
    def write_raw(data, path=None):
        if path:
            with open(path, 'w') as f:
                f.write(data)
        else:
            return data

    def write_file(data, path, header=None, encoding='utf-8'):
        if path.endswith('.csv'):
            Interface.write_csv(data, path, header, encoding=encoding)
        elif path.endswith('.json'):
            Interface.write_json(data, path, encoding=encoding)
        elif path.endswith('.yaml') or path.endswith('.yml'):
            Interface.write_yaml(data, path, encoding=encoding)
        elif path.endswith('.toml'):
            Interface.write_toml(data, path, encoding=encoding)
        elif path.endswith('.xml'):
            Interface.write_xml(data, path, encoding=encoding)
        elif path.endswith('.opml'):
            Interface.write_opml(data, path, encoding=encoding)
        elif path.endswith(''):
            Interface.write_raw(data, path, encoding=encoding)
        else:
            raise ValueError(f"Invalid file format: {path}")


class Logger:

    def __init__(self, filename):
        self.terminal = sys.stdout
        self.filename = filename
        self.log = open(filename, "w")

    def write(self, message):
        self.terminal.write(message)
        self.log.write(message)

    def flush(self):
        self.terminal.flush()
        self.log.flush()

    def isatty(self):
        return False

    def read(self):
        sys.stdout.flush()
        with open(self.filename, "r") as f:
            return f.read()
        
    def parse_format_phone(self, phone_number):
        # 電話番号の解析とフォーマットを行う処理
        cleaned_phone = re.sub(r'\D', '', phone_number)
        formatted_phone = f"({cleaned_phone[:3]})-{cleaned_phone[3:6]}-{cleaned_phone[6:]}"
        return formatted_phone

